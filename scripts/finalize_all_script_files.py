import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os
import shutil

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_documents():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    PRIMARY_COLOR = RGBColor(232, 130, 30)   # Má Hải Orange #E8821E
    SECONDARY_COLOR = RGBColor(0, 173, 239) # Má Hải Aqua #00ADEF
    DARK_COLOR = RGBColor(30, 41, 59)       # Slate-800 #1E293B
    GRAY_COLOR = RGBColor(100, 116, 139)   # Slate-500 #64748B
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = DARK_COLOR

    # --- HEADER / TITLE ---
    p_title_sub = doc.add_paragraph()
    p_title_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_title_sub.add_run("BÁNH MÌ MÁ HẢI · CỔNG ĐÀO TẠO & CHIA SẺ KINH NGHIỆM THỰC CHUYẾN")
    run_sub.font.size = Pt(11)
    run_sub.font.bold = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("KỊCH BẢN QUAY VIDEO & LỜI THOẠI AI VOICE\nHƯỚNG DẪN TRANG BỊ KINH NGHIỆM VẬN HÀNH DÀNH CHO ĐỐI TÁC NHƯỢNG QUYỀN")
    run_title.font.size = Pt(15)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR
    p_title.paragraph_format.space_after = Pt(12)

    # Info card box
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta.paragraph_format.space_after = Pt(16)
    
    r = p_meta.add_run("🌐 Website Đào Tạo: ")
    r.bold = True
    r = p_meta.add_run("https://daotao.banhmimahai.vn/nhuongquyen/\n")
    r.font.color.rgb = SECONDARY_COLOR
    
    r = p_meta.add_run("📱 Định dạng Video: ")
    r.bold = True
    r = p_meta.add_run("Mobile Dọc 9:16 (Quay màn hình smartphone)\n")
    
    r = p_meta.add_run("💡 Tông Giọng & Định Hướng: ")
    r.bold = True
    r = p_meta.add_run("Gần gũi, ấm áp, tập trung chia sẻ kinh nghiệm thực chiến từ các điểm bán thành công để đối tác tự tin vận hành đông khách (Không ký hiệu / /, không từ tâng bốc sáo rỗng)")

    # --- TABLE SECTION ---
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [Inches(1.8), Inches(2.5), Inches(3.7)]

    hdr_cells = table.rows[0].cells
    hdr_titles = ["Phân đoạn & Thời lượng", "Thao tác Quay màn hình Mobile", "Lời thoại AI Voice (Trang bị kinh nghiệm thực chiến)"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "E8821E")
        set_cell_margins(hdr_cells[i], top=160, bottom=160, left=180, right=180)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    trPr = table.rows[0]._element.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

    data = [
        (
            "PHẦN 1: Mở đầu & Giới thiệu Web\n(0:00 - 0:25)",
            [
                "Nhấn vào đường link: daotao.banhmimahai.vn/nhuongquyen/ (hoặc mở Safari/Chrome gõ link).",
                "Trang chủ xuất hiện Logo Bánh Mì Má Hải & Tiêu đề 'Đào Tạo Khách Hàng Nhượng Quyền'.",
                "Bấm icon [ ? ] (Hướng dẫn) ở góc trên bên phải header -> Mở Popup hướng dẫn -> Vuốt xem rồi đóng lại."
            ],
            "\"Xin chào Anh Chị đối tác nhượng quyền Bánh Mì Má Hải. Để giúp Anh Chị dễ dàng trang bị những kinh nghiệm thực chiến cần thiết từ các điểm bán thành công trong toàn hệ thống, Bánh Mì Má Hải đã tổng hợp trọn bộ kiến thức vận hành ngay trên cổng đào tạo trực tuyến này.\n\nAnh Chị có thể nhấn vào đường link được gửi, hoặc mở trình duyệt gõ daotao.banhmimahai.vn/nhuongquyen. Nếu muốn xem lại hướng dẫn nhanh bất cứ lúc nào, Anh Chị chỉ cần nhấn vào biểu tượng dấu chấm hỏi ở góc trên bên phải màn hình nhé.\""
        ),
        (
            "PHẦN 2: Nhập Thông tin Khách hàng\n(0:25 - 0:50)",
            [
                "Kéo xuống thẻ 'Thông tin khách hàng'.",
                "Ô Họ tên: Nhập chính xác tên mẫu Nguyễn Văn A.",
                "Ô Địa chỉ điểm bán: Nhập Xe 123 Nguyễn Thị Minh Khai, P.6, Q.3, TP.HCM."
            ],
            "\"Ngay tại màn hình chính, Anh Chị kéo xuống phần Thông tin khách hàng. Ô thứ nhất, Anh Chị điền Họ và tên của mình. Ô thứ hai, Anh Chị nhập Địa chỉ điểm bán hoặc tên xe bánh mì.\n\nĐiền đúng hai ô này sẽ giúp hệ thống lưu lại tiến độ học tập và đồng hành hỗ trợ Anh Chị trong suốt quá trình vận hành điểm bán.\""
        ),
        (
            "PHẦN 3: Giới thiệu 2 Chế độ trên Màn hình chính\n(0:50 - 1:20)",
            [
                "Kéo tiếp xuống khung Lựa chọn 2 Chế độ ngay bên dưới Form thông tin.",
                "Zoom chỉ tay vào card 'Chế độ Luyện tập' (màu cam).",
                "Zoom chỉ tay vào card 'Thi chính thức' (màu xanh Aqua)."
            ],
            "\"Ngay phía dưới phần thông tin, website có hai chế độ học tập rất linh hoạt.\n\nĐầu tiên là Chế độ Luyện tập, giúp Anh Chị thoải mái đúc kết kinh nghiệm theo từng chuyên đề, không bị giới hạn thời gian và có đáp án giải thích chi tiết sau mỗi câu hỏi.\n\nThứ hai là chế độ Kiểm tra kiến thức, giúp Anh Chị tự đánh giá lại mức độ nắm vững quy trình của mình qua 30 câu hỏi thực tế.\""
        ),
        (
            "PHẦN 4: Thao tác Chế độ Luyện tập (Ôn theo Chuyên đề)\n(1:20 - 2:05)",
            [
                "Nhấn nút 'Bắt đầu ngay' tại Chế độ Luyện tập.",
                "Màn hình mở ra Sidebar 'Lộ trình học tập' ở bên trái: Phần 1, Phần 2, Phần 3...",
                "Chạm chọn 1 Chuyên đề -> Chọn đáp án -> Hiện ngay thẻ Giải thích màu xanh -> Cuộn xuống xem thanh Tiến độ phần & Tiến độ lộ trình."
            ],
            "\"Bây giờ, chúng ta bấm vào nút Bắt đầu ngay ở ô Luyện tập. Màn hình sẽ mở ra trọn bộ bí quyết được đúc kết từ thực tế, bao gồm Vận hành điểm bán, Công thức làm bánh chuẩn vị, An toàn thực phẩm và Kỹ năng bán hàng thu hút khách.\n\nAnh Chị bấm chọn từng phần để rèn luyện, chạm chọn đáp án và xem lời giải thích chi tiết ngay bên dưới. Anh Chị cũng có thể cuộn xuống để theo dõi thanh tiến độ học tập của mình.\""
        ),
        (
            "PHẦN 5: Thao tác Chế độ Kiểm tra Rà soát Kiến thức\n(2:05 - 2:50)",
            [
                "Nhấn nút 'Thi chính thức' trên Navigation Header.",
                "Màn hình hiển thị đếm ngược 30 phút & 30 câu hỏi ngẫu nhiên.",
                "Chọn đáp án -> Bấm 'Câu kế tiếp'.",
                "Chạm vào 'Bảng câu hỏi' (ô màu xanh Aqua là câu đã làm) -> Kéo xuống bấm 'Nộp bài thi' -> Xác nhận."
            ],
            "\"Khi đã nắm vững các kinh nghiệm, Anh Chị bấm sang mục Kiểm tra kiến thức trên thanh menu phía trên. Bộ đề gồm 30 câu hỏi thực tế giúp Anh Chị rà soát lại toàn bộ quy trình trước khi mở bán.\n\nAnh Chị chọn đáp án đúng rồi bấm Câu kế tiếp, hoặc bấm Bảng câu hỏi để xem lại các câu đã chọn. Khi hoàn thành, Anh Chị kéo xuống bấm Nộp bài thi và chọn Xác nhận.\""
        ),
        (
            "PHẦN 6: Xem Kết quả & Tải Chứng Nhận Đồng Hành (PNG)\n(2:50 - 3:20)",
            [
                "Hiện màn hình Kết quả tích xanh 'VƯỢT QUA: 26/30'.",
                "Vuốt xuống xem khung CHỨNG NHẬN ĐÀO TẠO sang trọng in tên Nguyễn Văn A & Địa chỉ điểm bán.",
                "Nhấn nút màu cam 'Tải chứng nhận (PNG)' -> Ảnh tự động lưu về máy."
            ],
            "\"Vừa nộp bài xong, màn hình sẽ hiển thị kết quả kiểm tra. Đạt từ 20 trên 30 câu đúng là Anh Chị đã hoàn toàn tự tin để vận hành điểm bán.\n\nKhung Chứng nhận hoàn thành đào tạo của Bánh Mì Má Hải sẽ hiện ra như một lời ghi nhận cho sự chuẩn bị chu đáo của Anh Chị. Anh Chị nhấn vào nút màu cam Tải chứng nhận PNG để lưu kỷ niệm bằng khen về điện thoại của mình nhé.\""
        ),
        (
            "PHẦN 7: Outro / Đồng hành cùng Má Hải\n(3:20 - 3:35)",
            [
                "Màn hình kết thúc xuất hiện Logo Bánh Mì Má Hải & Hotline hỗ trợ đào tạo.",
                "Slogan thương hiệu & Thông tin liên hệ phòng Đào tạo."
            ],
            "\"Hy vọng Cổng đào tạo nhượng quyền Bánh Mì Má Hải sẽ là người bạn đồng hành đắc lực, giúp Anh Chị làm chủ quy trình và tự tin mở bán đông khách mỗi ngày. Chúc Anh Chị kinh doanh hồng phát và bùng nổ doanh số cùng Bánh Mì Má Hải. Xin chào và hẹn gặp lại.\""
        )
    ]

    for row_idx, (section_title, steps, voiceover) in enumerate(data):
        row_cells = table.add_row().cells
        bg_color = "FAF7F2" if row_idx % 2 == 0 else "FFFFFF"
        
        for i in range(3):
            row_cells[i].width = col_widths[i]
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=140, bottom=140, left=160, right=160)
            
        p0 = row_cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_sec = p0.add_run(section_title)
        run_sec.bold = True
        run_sec.font.size = Pt(10)
        run_sec.font.color.rgb = DARK_COLOR
        
        p1 = row_cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for idx, step in enumerate(steps):
            if idx > 0:
                p1 = row_cells[1].add_paragraph()
            p1.paragraph_format.space_after = Pt(4)
            r_bullet = p1.add_run("• ")
            r_bullet.bold = True
            r_bullet.font.color.rgb = PRIMARY_COLOR
            r_text = p1.add_run(step)
            r_text.font.size = Pt(9.5)

        p2 = row_cells[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_voice = p2.add_run(voiceover)
        r_voice.italic = True
        r_voice.font.size = Pt(10)
        r_voice.font.color.rgb = DARK_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    p_tips_head = doc.add_paragraph()
    r_tips_head = p_tips_head.add_run("💡 ĐỊNH HƯỚNG GIỌNG ĐỌC & TÂM LÝ TRUYỀN TẢI (EMPOWERING TONE)")
    r_tips_head.bold = True
    r_tips_head.font.size = Pt(12)
    r_tips_head.font.color.rgb = PRIMARY_COLOR
    p_tips_head.paragraph_format.space_after = Pt(6)

    tips = [
        ("1. Tông giọng chia sẻ kinh nghiệm:", "Giọng đọc ấm áp, chân thành, mang tính đồng hành động viên (không giống giám thị coi thi hay áp đặt chỉ tiêu)."),
        ("2. Nhấn mạnh giá trị thực chiến:", "Tập trung vào các từ khóa: trang bị kinh nghiệm, công thức làm bánh chuẩn vị, vận hành đông khách, bùng nổ doanh số."),
        ("3. Tốc độ đọc AI Voice:", "Giữ nhịp thong thả (0.95x - 1.0x), để khoảng dừng vừa phải sau mỗi ý để người xem thấm bài."),
        ("4. Bằng khen Chứng nhận:", "Nhấn mạnh Bằng khen là 'sự ghi nhận chu đáo' và 'kỷ niệm đồng hành' chứ không tạo áp lực thi cử nặng nề.")
    ]

    for t_title, t_desc in tips:
        p_tip = doc.add_paragraph()
        p_tip.paragraph_format.space_after = Pt(4)
        p_tip.paragraph_format.left_indent = Inches(0.2)
        r_t = p_tip.add_run(t_title + " ")
        r_t.bold = True
        r_t.font.color.rgb = SECONDARY_COLOR
        r_d = p_tip.add_run(t_desc)
        r_d.font.size = Pt(10)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(18)
    r_f = p_foot.add_run("--- BÁNH MÌ MÁ HẢI · PHÒNG ĐÀO TẠO & CHIA SẺ KINH NGHIỆM THỰC CHUYẾN ---")
    r_f.font.size = Pt(9)
    r_f.font.color.rgb = GRAY_COLOR

    # Outputs folder
    out_docx = r"outputs\Kich_Ban_Video_Huong_Dan_Nhuong_Quyen_Ban_Mi_Ma_Hai.docx"
    desktop_docx = r"C:\Users\admin\Desktop\KICH_BAN_QUAY_VIDEO_VA_THU_AM_VOICEOVER_MA_HAI.docx"
    
    doc.save(out_docx)
    shutil.copy(out_docx, desktop_docx)
    print(f"SUCCESS: Saved to {out_docx} and {desktop_docx}")

if __name__ == "__main__":
    create_documents()
