import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # Page Margins: 1 inch (72 pt / 1440 dxa)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles
    PRIMARY_COLOR = RGBColor(232, 130, 30)   # Má Hải Orange #E8821E
    SECONDARY_COLOR = RGBColor(0, 173, 239) # Má Hải Aqua #00ADEF
    DARK_COLOR = RGBColor(30, 41, 59)       # Slate-800 #1E293B
    GRAY_COLOR = RGBColor(100, 116, 139)   # Slate-500 #64748B
    
    # Set default font
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = DARK_COLOR

    # --- HEADER / TITLE SECTION ---
    p_title_sub = doc.add_paragraph()
    p_title_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_title_sub.add_run("BÁNH MÌ MÁ HẢI · HỆ THỐNG ĐÀO TẠO NHƯỢNG QUYỀN")
    run_sub.font.size = Pt(11)
    run_sub.font.bold = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("KỊCH BẢN QUAY VIDEO & THU ÂM VOICE-OVER\nHƯỚNG DẪN SỬ DỤNG WEBSITE DÀNH CHO KHÁCH HÀNG NHƯỢNG QUYỀN")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR
    p_title.paragraph_format.space_after = Pt(12)

    # Info card box
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta.paragraph_format.space_after = Pt(18)
    
    r = p_meta.add_run("📌 Link Cổng Đào Tạo: ")
    r.bold = True
    r = p_meta.add_run("https://daotao.banhmimahai.vn/nhuongquyen/\n")
    r.font.color.rgb = SECONDARY_COLOR
    
    r = p_meta.add_run("📱 Định dạng Video: ")
    r.bold = True
    r = p_meta.add_run("Mobile Dọc 9:16 (Quay màn hình smartphone)\n")
    
    r = p_meta.add_run("⏱️ Thời lượng dự kiến: ")
    r.bold = True
    r = p_meta.add_run("3 phút 35 giây (Cấu trúc 7 phần chuẩn kết cấu UI Web)")

    # --- TABLE SECTION ---
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Widths
    col_widths = [Inches(1.8), Inches(2.6), Inches(3.6)]

    # Table Header Row
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Phân đoạn & Thời lượng", "Thao tác Quay màn hình Mobile", "Lời thoại Voiceover (Có ngắt nghỉ `/`)"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "E8821E")
        set_cell_margins(hdr_cells[i], top=160, bottom=160, left=180, right=180)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    # Repeat header row on every page
    trPr = table.rows[0]._element.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

    # Script Data
    data = [
        (
            "PHẦN 1: Mở đầu & Giới thiệu Web\n(0:00 - 0:25)",
            [
                "Nhấn vào đường link gửi sẵn: daotao.banhmimahai.vn/nhuongquyen/ (hoặc mở Safari/Chrome gõ link).",
                "Trang chủ xuất hiện Logo Bánh Mì Má Hải & Tiêu đề 'Đào Tạo Khách Hàng Nhượng Quyền'.",
                "Bấm icon [ ? ] (Hướng dẫn) ở góc trên bên phải header -> Mở Popup hướng dẫn -> Vuốt xem rồi đóng lại."
            ],
            "\"Xin chào tất cả Anh/Chị đối tác nhượng quyền / của Bánh Mì Má Hải!\n\nĐể giúp Anh/Chị nắm vững quy trình vận hành điểm bán / và dễ dàng hoàn thành kỳ kiểm tra cấp chứng nhận / Bánh Mì Má Hải đã xây dựng Cổng đào tạo trực tuyến ngay trên điện thoại.\n\nAnh/Chị có thể nhấn thẳng vào đường link được gửi / hoặc mở trình duyệt nhập daotao.banhmimahai.vn/nhuongquyen.\n\nNếu cần xem lại hướng dẫn nhanh bất kỳ lúc nào / Anh/Chị nhấn vào biểu tượng dấu chấm hỏi ở góc trên cùng bên phải màn hình nhé!\""
        ),
        (
            "PHẦN 2: Nhập Thông tin Khách hàng\n(0:25 - 0:50)",
            [
                "Kéo xuống thẻ 'Thông tin khách hàng'.",
                "Ô Họ tên: Nhập tên mẫu Nguyễn Văn A.",
                "Ô Địa chỉ điểm bán: Nhập Xe 123 Nguyễn Thị Minh Khai, P.6, Q.3, TP.HCM."
            ],
            "\"Ngay tại màn hình chính / bước đầu tiên vô cùng quan trọng mà Anh/Chị cần lưu ý / đó là điền đầy đủ Thông tin khách hàng.\n\nTại ô thứ nhất / Anh/Chị nhập chính xác Họ và tên của mình.\n\nTại ô thứ hai / Anh/Chị nhập Địa chỉ hoặc tên xe / điểm bán của Anh/Chị.\n\nThông tin này rất quan trọng / vì hệ thống sẽ tự động in Họ tên và Địa chỉ này / trực tiếp lên Bằng khen Chứng nhận đào tạo chính thức của Bánh Mì Má Hải khi Anh/Chị hoàn thành bài thi đấy ạ!\""
        ),
        (
            "PHẦN 3: Giới thiệu 2 Chế độ trên Màn hình chính\n(0:50 - 1:20)",
            [
                "Kéo tiếp xuống khung Lựa chọn 2 Chế độ ngay bên dưới Form thông tin.",
                "Zoom chỉ tay vào card 'Chế độ Luyện tập' (màu cam).",
                "Zoom chỉ tay vào card 'Thi chính thức' (màu xanh Aqua)."
            ],
            "\"Ngay phía dưới phần thông tin / hệ thống mang đến 2 lựa chọn phù hợp với nhu cầu của Anh/Chị:\n\nChế độ thứ nhất là 'Chế độ Luyện tập' / dành cho Anh/Chị mới bắt đầu ôn lại kiến thức / học theo từng chuyên đề / không bị giới hạn thời gian và có đáp án giải thích ngay.\n\nChế độ thứ hai là 'Thi chính thức' / dành cho Anh/Chị đã tự tin / chuẩn bị làm bài thi 30 câu trong 30 phút để nhận Chứng nhận.\""
        ),
        (
            "PHẦN 4: Thao tác Chế độ Luyện tập (Ôn theo Chuyên đề)\n(1:20 - 2:05)",
            [
                "Nhấn nút 'Bắt đầu ngay' tại Chế độ Luyện tập.",
                "Màn hình mở ra Sidebar 'Lộ trình học tập' ở bên trái (vuốt ngang trên Mobile): Phần 1, Phần 2, Phần 3...",
                "Chạm chọn 1 Chuyên đề -> Chọn đáp án -> Hiện ngay thẻ Giải thích màu xanh -> Cuộn xuống xem thanh Tiến độ phần & Tiến độ lộ trình."
            ],
            "\"Bây giờ / chúng ta bấm vào 'Bắt đầu ngay' tại Chế độ Luyện tập.\n\nGiao diện sẽ hiển thị 'Lộ trình học tập' với các chuyên đề được chia nhỏ vô cùng khoa học / như Vận hành, Công thức, ATVSTP và Tiếp thị.\n\nAnh/Chị bấm chọn từng phần để ôn / chạm chọn đáp án / hệ thống sẽ báo Đúng hay Chưa chính xác kèm lời giải thích chi tiết ngay bên dưới.\n\nAnh/Chị cũng có thể theo dõi thanh phần trăm Tiến độ lộ trình ở cuối màn hình đấy ạ!\""
        ),
        (
            "PHẦN 5: Thao tác Chế độ Thi chính thức & Nộp bài\n(2:05 - 2:50)",
            [
                "Nhấn nút 'Thi chính thức' trên Navigation Header.",
                "Màn hình hiển thị đếm ngược 30 phút & 30 câu hỏi ngẫu nhiên.",
                "Chọn đáp án -> Bấm 'Câu kế tiếp'.",
                "Chạm vào 'Bảng câu hỏi' (ô màu xanh Aqua là câu đã làm) -> Kéo xuống bấm 'Nộp bài thi' -> Xác nhận."
            ],
            "\"Khi đã thuộc bài / Anh/Chị bấm chuyển sang tab 'Thi chính thức' trên thanh menu trên cùng.\n\nBộ đề thi sẽ gồm 30 câu hỏi ngẫu nhiên với thời gian làm bài là 30 phút.\n\nAnh/Chị chỉ cần chọn đáp án và bấm 'Câu kế tiếp'.\n\nCó thể bấm vào 'Bảng câu hỏi' để xem tổng quan các câu đã trả lời / và bấm 'Nộp bài thi' khi hoàn thành 30 câu.\""
        ),
        (
            "PHẦN 6: Xem Kết quả & Tải Chứng Nhận (PNG)\n(2:50 - 3:20)",
            [
                "Hiện màn hình Kết quả tích xanh 'VƯỢT QUA: 26/30'.",
                "Vuốt xuống xem khung CHỨNG NHẬN ĐÀO TẠO sang trọng in tên Nguyễn Văn A & Địa chỉ điểm bán.",
                "Nhấn nút màu cam 'Tải chứng nhận (PNG)' -> Ảnh tự động lưu về máy."
            ],
            "\"Ngay khi nộp bài / màn hình sẽ trả kết quả lập tức.\n\nChỉ cần đúng từ 20 trên 30 câu trở lên / Anh/Chị sẽ nhận thông báo VƯỢT QUA!\n\nKhung 'CHỨNG NHẬN ĐÀO TẠO' chính thức của Bánh Mì Má Hải sẽ hiện ra với đầy đủ Họ tên, Địa chỉ điểm bán và Mã chứng nhận riêng.\n\nAnh/Chị chỉ cần nhấn nút màu cam 'Tải chứng nhận (PNG)' / bức ảnh bằng khen nét cao sẽ được tải trực tiếp về thư viện ảnh trên điện thoại!\""
        ),
        (
            "PHẦN 7: Outro / Kết bài\n(3:20 - 3:35)",
            [
                "Màn hình kết thúc xuất hiện Logo Bánh Mì Má Hải & Hotline hỗ trợ đào tạo.",
                "Slogan thương hiệu & Thông tin liên hệ phòng Đào tạo."
            ],
            "\"Cổng đào tạo nhượng quyền Bánh Mì Má Hải thật đơn giản và chuyên nghiệp đúng không ạ?\n\nChúc Anh/Chị đối tác ôn tập thật tốt / đạt kết quả thật cao và kinh doanh phát đạt cùng Bánh Mì Má Hải!\n\nXin chào và hẹn gặp lại!\""
        )
    ]

    for row_idx, (section_title, steps, voiceover) in enumerate(data):
        row_cells = table.add_row().cells
        bg_color = "FAF7F2" if row_idx % 2 == 0 else "FFFFFF"
        
        for i in range(3):
            row_cells[i].width = col_widths[i]
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=140, bottom=140, left=160, right=160)
            
        # Col 0: Section Title & Duration
        p0 = row_cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_sec = p0.add_run(section_title)
        run_sec.bold = True
        run_sec.font.size = Pt(10)
        run_sec.font.color.rgb = DARK_COLOR
        
        # Col 1: Steps (Mobile actions)
        p1 = row_cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for idx, step in enumerate(steps):
            if idx > 0:
                p1 = row_cells[1].add_paragraph()
            p1.paragraph_format.space_after = Pt(4)
            r_bullet = p1.add_run("• ")
            r_bullet.bold = True
            r_bullet.font.color.rgb = PRIMARY_COLOR
            r_text = p1.add_run(step)
            r_text.font.size = Pt(9.5)

        # Col 2: Voiceover Script
        p2 = row_cells[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_voice = p2.add_run(voiceover)
        r_voice.italic = True
        r_voice.font.size = Pt(10)
        r_voice.font.color.rgb = DARK_COLOR

    # --- BEST PRACTICES SECTION ---
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    p_tips_head = doc.add_paragraph()
    r_tips_head = p_tips_head.add_run("💡 4 MẸO KỸ THUẬT QUAY & DỰNG VIDEO ĐẠT HIỆU QUẢ CAO NHẤT")
    r_tips_head.bold = True
    r_tips_head.font.size = Pt(12)
    r_tips_head.font.color.rgb = PRIMARY_COLOR
    p_tips_head.paragraph_format.space_after = Pt(6)

    tips = [
        ("1. Kích hoạt điểm chạm màn hình (Touch Pointer):", "Bật chế độ hiển thị điểm chạm tay trên Android/iOS hoặc chèn hiệu ứng vòng tròn chạm (Circle Pulse) trong CapCut/Premiere khi bấm vào nút 'Bảng câu hỏi', 'Nộp bài thi', 'Tải chứng nhận (PNG)' để người xem dễ thao tác theo."),
        ("2. Tỷ lệ khung hình 9:16 (Màn hình dọc):", "Dùng điện thoại thông minh để quay toàn bộ thao tác màn hình dọc, tối ưu cho việc gửi Zalo, Facebook Messenger hoặc đăng TikTok/Reels/Shorts."),
        ("3. Zoom 1.2x – 1.5x vào Bằng khen PNG:", "Tại Phần 6, zoom cận cảnh thông tin Họ tên + Địa chỉ xe nhượng quyền + Con dấu Bánh Mì Má Hải trên mẫu chứng nhận để tạo niềm tin và sự hào hứng cho đối tác."),
        ("4. Nhạc nền (Background Music):", "Chọn nhạc tươi vui, truyền cảm hứng nhưng giữ âm lượng ở mức 10% - 15% để tôn giọng thu âm Voiceover rõ ràng.")
    ]

    for t_title, t_desc in tips:
        p_tip = doc.add_paragraph()
        p_tip.paragraph_format.space_after = Pt(4)
        p_tip.paragraph_format.left_indent = Inches(0.2)
        r_t = p_tip.add_run(t_title + " ")
        r_t.bold = True
        r_t.font.color.rgb = SECONDARY_COLOR
        r_d = p_tip.add_run(t_desc)
        r_d.font.size = Pt(10)

    # Footer note
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(18)
    r_f = p_foot.add_run("--- BÁNH MÌ MÁ HẢI · PHÒNG ĐÀO TẠO & VẬN HÀNH NHƯỢNG QUYỀN ---")
    r_f.font.size = Pt(9)
    r_f.font.color.rgb = GRAY_COLOR

    # Ensure output folder exists
    output_dir = "outputs"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    output_path = os.path.join(output_dir, "Kich_Ban_Video_Huong_Dan_Nhuong_Quyen_Ban_Mi_Ma_Hai.docx")
    doc.save(output_path)
    print(f"SUCCESS: Document saved to {output_path}")

if __name__ == "__main__":
    create_document()
